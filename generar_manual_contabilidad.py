from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ─── Estilos ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xA0)

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# ─── Funciones auxiliares ───
def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_paragraph(text, bold=False, style_name=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_step(number, title, description):
    p = doc.add_paragraph()
    run_num = p.add_run(f'Paso {number}: ')
    run_num.bold = True
    run_num.font.color.rgb = RGBColor(0x1A, 0x56, 0xA0)
    run_title = p.add_run(title)
    run_title.bold = True
    if description:
        doc.add_paragraph(description)

def add_note(text):
    p = doc.add_paragraph()
    run_icon = p.add_run('NOTA: ')
    run_icon.bold = True
    run_icon.font.color.rgb = RGBColor(0xD4, 0x6B, 0x08)
    p.add_run(text)

def add_tip(text):
    p = doc.add_paragraph()
    run_icon = p.add_run('CONSEJO: ')
    run_icon.bold = True
    run_icon.font.color.rgb = RGBColor(0x16, 0x7F, 0x39)
    p.add_run(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MANUAL DE USUARIO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xA0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Módulo de Contabilidad')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

system = doc.add_paragraph()
system.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = system.add_run('Sistema de Facturación CompuServices')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Febrero 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Versión 1.0')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS
# ══════════════════════════════════════════════════════════
doc.add_heading('Tabla de Contenidos', level=1)

toc_items = [
    '1. Introducción',
    '2. Requisitos y Acceso',
    '3. Plan de Cuentas',
    '   3.1 Ver el plan de cuentas',
    '   3.2 Crear una nueva cuenta',
    '   3.3 Editar una cuenta',
    '   3.4 Eliminar una cuenta',
    '4. Asientos Contables',
    '   4.1 Ver listado de asientos',
    '   4.2 Filtrar asientos',
    '   4.3 Crear un asiento manual',
    '   4.4 Ver detalle de un asiento',
    '   4.5 Editar un asiento (borrador)',
    '   4.6 Contabilizar un asiento',
    '   4.7 Anular un asiento',
    '5. Generación Automática de Asientos',
    '   5.1 Desde una factura',
    '   5.2 Desde un gasto',
    '   5.3 Desde un retiro de caja',
    '6. Reportes Contables',
    '   6.1 Libro Diario',
    '   6.2 Libro Mayor',
    '   6.3 Balance de Comprobación',
    '   6.4 Balance General',
    '   6.5 Estado de Resultados',
    '7. Preguntas Frecuentes',
    '8. Glosario',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', level=1)

doc.add_paragraph(
    'El módulo de Contabilidad del Sistema de Facturación CompuServices permite llevar un registro '
    'ordenado y detallado de todas las transacciones financieras del negocio mediante asientos contables. '
    'Este módulo se integra directamente con los módulos de Facturación, Gastos y Retiros para generar '
    'asientos de forma automática.'
)

doc.add_paragraph('Con este módulo usted podrá:')

features = [
    'Administrar el Plan de Cuentas contable de su empresa.',
    'Registrar asientos contables manuales con múltiples líneas de detalle.',
    'Generar asientos automáticamente desde facturas, gastos y retiros.',
    'Contabilizar o anular asientos según el flujo de aprobación.',
    'Consultar reportes contables: Libro Diario, Libro Mayor, Balance de Comprobación, Balance General y Estado de Resultados.',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 2. REQUISITOS Y ACCESO
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Requisitos y Acceso', level=1)

doc.add_heading('¿Quién puede acceder?', level=2)
doc.add_paragraph(
    'El módulo de Contabilidad está disponible únicamente para los siguientes tipos de usuario:'
)

add_table(
    ['Tipo de Usuario', 'Acceso al Módulo', 'Permisos'],
    [
        ['Administrador', 'Completo', 'Todas las funciones (crear, editar, eliminar, contabilizar, anular, reportes)'],
        ['Super Usuario', 'Completo', 'Todas las funciones'],
        ['Técnico', 'Sin acceso', 'No puede ver ni acceder al módulo'],
        ['Atención al Público', 'Sin acceso', 'No puede ver ni acceder al módulo'],
    ]
)

doc.add_paragraph()

doc.add_heading('¿Cómo acceder al módulo?', level=2)

add_step(1, 'Iniciar sesión', 'Ingrese al sistema con su usuario y contraseña de Administrador o Super Usuario.')
add_step(2, 'Navegar al módulo', 'En el menú lateral izquierdo, busque y haga clic en "Contabilidad".')
add_step(3, 'Seleccionar sección', 'El módulo se divide en 3 pestañas principales: Plan de Cuentas, Asientos Contables y Reportes.')

add_note('Si no ve la opción "Contabilidad" en el menú, verifique que su usuario tenga el tipo Administrador o Super Usuario.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 3. PLAN DE CUENTAS
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Plan de Cuentas', level=1)

doc.add_paragraph(
    'El Plan de Cuentas es la estructura contable base de su empresa. Define todas las cuentas donde se '
    'registrarán los movimientos financieros. El sistema viene con un plan de cuentas predeterminado '
    'adaptado para negocios de servicios técnicos y venta de productos en Ecuador.'
)

# 3.1
doc.add_heading('3.1 Ver el Plan de Cuentas', level=2)

add_step(1, 'Acceder', 'Ingrese al módulo de Contabilidad y seleccione la pestaña "Plan de Cuentas".')
add_step(2, 'Explorar', 'Verá el plan de cuentas en formato de árbol jerárquico con 5 grupos principales:')

add_table(
    ['Código', 'Grupo', 'Descripción'],
    [
        ['1', 'ACTIVOS', 'Lo que la empresa posee (caja, bancos, inventario, cuentas por cobrar)'],
        ['2', 'PASIVOS', 'Lo que la empresa debe (cuentas por pagar, IVA, préstamos)'],
        ['3', 'PATRIMONIO', 'Capital y resultados acumulados de la empresa'],
        ['4', 'INGRESOS', 'Dinero que entra por ventas y servicios'],
        ['5', 'GASTOS', 'Dinero que sale por operaciones del negocio'],
    ]
)

doc.add_paragraph()
add_step(3, 'Expandir/Colapsar', 'Haga clic en la flecha junto a cada grupo para ver sus subcuentas. Las cuentas marcadas como "detalle" son las que reciben movimientos contables.')

doc.add_paragraph()
doc.add_paragraph('El plan de cuentas predeterminado incluye las siguientes cuentas de detalle:')

add_table(
    ['Código', 'Cuenta', 'Tipo', 'Uso Principal'],
    [
        ['1.1.01', 'Caja', 'Activo', 'Dinero en efectivo y cobros'],
        ['1.1.02', 'Bancos', 'Activo', 'Dinero en cuentas bancarias'],
        ['1.1.03', 'Cuentas por Cobrar', 'Activo', 'Ventas a crédito pendientes'],
        ['1.1.04', 'Inventario de Mercaderías', 'Activo', 'Valor del stock de productos'],
        ['1.1.05', 'IVA en Compras', 'Activo', 'Crédito tributario por IVA'],
        ['1.2.01', 'Equipos y Maquinaria', 'Activo', 'Activos fijos del negocio'],
        ['1.2.02', 'Muebles y Enseres', 'Activo', 'Mobiliario del negocio'],
        ['1.2.03', '(-) Depreciación Acumulada', 'Activo', 'Desgaste de activos fijos'],
        ['2.1.01', 'Cuentas por Pagar', 'Pasivo', 'Deudas con proveedores'],
        ['2.1.02', 'IVA en Ventas', 'Pasivo', 'IVA cobrado pendiente de declarar'],
        ['2.1.03', 'Retenciones por Pagar', 'Pasivo', 'Retenciones pendientes'],
        ['2.1.04', 'Sueldos por Pagar', 'Pasivo', 'Nómina pendiente'],
        ['2.2.01', 'Préstamos Bancarios', 'Pasivo', 'Deudas con bancos'],
        ['3.1.01', 'Capital Social', 'Patrimonio', 'Inversión inicial del negocio'],
        ['3.1.02', 'Resultados del Ejercicio', 'Patrimonio', 'Utilidad/pérdida del año actual'],
        ['3.1.03', 'Resultados Acumulados', 'Patrimonio', 'Utilidades de años anteriores'],
        ['4.1.01', 'Ventas de Productos', 'Ingreso', 'Ingresos por venta de mercadería'],
        ['4.1.02', 'Ingresos por Servicios Técnicos', 'Ingreso', 'Ingresos por reparaciones'],
        ['4.2.01', 'Otros Ingresos', 'Ingreso', 'Ingresos no operacionales'],
        ['5.1.01', 'Costo de Ventas', 'Gasto', 'Costo de la mercadería vendida'],
        ['5.1.02', 'Sueldos y Salarios', 'Gasto', 'Pago de nómina'],
        ['5.1.03', 'Servicios Básicos', 'Gasto', 'Luz, agua, teléfono, internet'],
        ['5.1.04', 'Arriendo', 'Gasto', 'Pago de alquiler del local'],
        ['5.1.05', 'Suministros de Oficina', 'Gasto', 'Material de oficina'],
        ['5.1.06', 'Depreciaciones', 'Gasto', 'Gasto por desgaste de activos'],
        ['5.1.07', 'Gastos Generales', 'Gasto', 'Otros gastos operacionales'],
        ['5.2.01', 'Gastos Financieros', 'Gasto', 'Intereses y comisiones bancarias'],
        ['5.2.02', 'Otros Gastos', 'Gasto', 'Gastos no operacionales'],
    ]
)

doc.add_paragraph()

# 3.2
doc.add_heading('3.2 Crear una nueva cuenta', level=2)

doc.add_paragraph('Si necesita agregar una cuenta que no existe en el plan predeterminado:')

add_step(1, 'Abrir formulario', 'Haga clic en el botón "Nueva Cuenta" ubicado en la parte superior derecha.')
add_step(2, 'Completar los campos', '')

add_table(
    ['Campo', 'Descripción', 'Ejemplo'],
    [
        ['Código', 'Código único jerárquico de la cuenta', '5.1.08'],
        ['Nombre', 'Nombre descriptivo de la cuenta', 'Publicidad y Marketing'],
        ['Tipo', 'Clasificación contable', 'Gasto'],
        ['Naturaleza', 'Deudora (activos/gastos) o Acreedora (pasivos/patrimonio/ingresos)', 'Deudora'],
        ['Cuenta Padre', 'Cuenta superior en la jerarquía', '5.1 - Gastos Operacionales'],
        ['Nivel', 'Profundidad en el árbol (1=grupo, 2=subgrupo, 3=detalle)', '3'],
        ['Es Detalle', 'Marcar si la cuenta recibirá movimientos contables', 'Sí'],
    ]
)

doc.add_paragraph()
add_step(3, 'Guardar', 'Haga clic en "Guardar". La nueva cuenta aparecerá en el árbol.')

add_tip(
    'Si la cuenta es de tipo Activo o Gasto, su naturaleza debe ser "Deudora". '
    'Si es Pasivo, Patrimonio o Ingreso, su naturaleza debe ser "Acreedora".'
)

# 3.3
doc.add_heading('3.3 Editar una cuenta', level=2)
add_step(1, 'Ubicar la cuenta', 'Navegue en el árbol hasta encontrar la cuenta que desea modificar.')
add_step(2, 'Editar', 'Haga clic en el botón "Editar" junto a la cuenta.')
add_step(3, 'Modificar', 'Cambie los campos necesarios y haga clic en "Guardar".')

add_note('Se recomienda no cambiar el código de una cuenta que ya tiene movimientos registrados.')

# 3.4
doc.add_heading('3.4 Eliminar una cuenta', level=2)
add_step(1, 'Ubicar la cuenta', 'Navegue en el árbol hasta la cuenta a eliminar.')
add_step(2, 'Eliminar', 'Haga clic en el botón "Eliminar" y confirme la acción.')

add_note('No se puede eliminar una cuenta que tenga movimientos contables registrados ni una cuenta que tenga subcuentas. Primero debe eliminar o reasignar los movimientos y subcuentas.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 4. ASIENTOS CONTABLES
# ══════════════════════════════════════════════════════════
doc.add_heading('4. Asientos Contables', level=1)

doc.add_paragraph(
    'Un asiento contable es el registro de una transacción financiera. Cada asiento tiene una cabecera '
    '(fecha, descripción, tipo) y múltiples líneas de detalle donde se indica qué cuentas se afectan '
    'y por qué monto. La regla fundamental es que la suma del Debe siempre debe ser igual a la suma '
    'del Haber (principio de partida doble).'
)

doc.add_heading('Estados de un asiento', level=3)

add_table(
    ['Estado', 'Color', 'Significado', '¿Se puede editar?'],
    [
        ['Borrador', 'Amarillo', 'Recién creado, pendiente de revisión y aprobación', 'Sí'],
        ['Contabilizado', 'Verde', 'Aprobado. Afecta los saldos y aparece en reportes', 'No'],
        ['Anulado', 'Rojo', 'Cancelado. No afecta saldos ni reportes', 'No'],
    ]
)

doc.add_paragraph()

doc.add_heading('Tipos de asiento', level=3)

add_table(
    ['Tipo', 'Origen', 'Descripción'],
    [
        ['Manual', 'Creado por el usuario', 'Asiento ingresado manualmente desde el módulo'],
        ['Venta', 'Factura (contado)', 'Generado automáticamente al contabilizar una factura de contado'],
        ['Crédito', 'Factura (crédito)', 'Generado desde una factura de venta a crédito'],
        ['Abono Crédito', 'Abono a crédito', 'Generado al registrar un abono a un crédito'],
        ['Gasto', 'Módulo de Gastos', 'Generado desde un gasto registrado'],
        ['Retiro', 'Retiro de caja', 'Generado desde un retiro de caja'],
        ['Anulación', 'Anulación de factura', 'Reversa el asiento de una factura anulada'],
        ['Ajuste', 'Manual', 'Asiento de ajuste contable'],
        ['Cierre', 'Manual', 'Asiento de cierre de periodo contable'],
    ]
)

doc.add_paragraph()

# 4.1
doc.add_heading('4.1 Ver listado de asientos', level=2)
add_step(1, 'Acceder', 'En el módulo de Contabilidad, seleccione la pestaña "Asientos Contables".')
add_step(2, 'Revisar', 'Verá una tabla con todos los asientos ordenados del más reciente al más antiguo. Cada fila muestra: número, fecha, descripción, tipo, estado, total debe y total haber.')

# 4.2
doc.add_heading('4.2 Filtrar asientos', level=2)
doc.add_paragraph('Utilice los filtros en la parte superior de la tabla:')
add_table(
    ['Filtro', 'Descripción'],
    [
        ['Fecha Desde / Hasta', 'Rango de fechas para buscar asientos'],
        ['Tipo', 'Filtrar por tipo: Manual, Venta, Gasto, Retiro, etc.'],
        ['Estado', 'Filtrar por estado: Borrador, Contabilizado, Anulado'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('Después de seleccionar los filtros, haga clic en "Filtrar" para aplicar la búsqueda.')

# 4.3
doc.add_heading('4.3 Crear un asiento manual', level=2)

doc.add_paragraph('Para registrar una transacción que no proviene de facturación, gastos o retiros:')

add_step(1, 'Nuevo asiento', 'Haga clic en el botón "Nuevo Asiento".')
add_step(2, 'Completar cabecera', '')

add_table(
    ['Campo', 'Descripción', 'Ejemplo'],
    [
        ['Fecha', 'Fecha de la transacción', '15/02/2026'],
        ['Descripción', 'Detalle de la operación', 'Pago de arriendo del local - Febrero 2026'],
    ]
)

doc.add_paragraph()
add_step(3, 'Agregar líneas de detalle', 'Complete la tabla de líneas con las cuentas afectadas:')

add_table(
    ['Cuenta', 'Descripción', 'Debe', 'Haber'],
    [
        ['5.1.04 - Arriendo', 'Arriendo febrero', '$500.00', ''],
        ['1.1.01 - Caja', 'Pago desde caja', '', '$500.00'],
        ['TOTALES', '', '$500.00', '$500.00'],
    ]
)

doc.add_paragraph()

doc.add_paragraph('Reglas para las líneas:')
items = [
    'Mínimo 2 líneas por asiento.',
    'Cada línea debe tener un monto en Debe O en Haber, nunca en ambos.',
    'El total del Debe debe ser exactamente igual al total del Haber.',
    'Use el botón "+ Agregar línea" para añadir más filas.',
    'Use el botón "X" para eliminar una línea.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_step(4, 'Verificar', 'El sistema mostrará si el asiento está "Cuadrado" (correcto) o tiene una diferencia. Solo podrá guardar si está cuadrado.')
add_step(5, 'Guardar', 'Haga clic en "Guardar". El asiento se creará en estado "Borrador".')

add_note('Un asiento en estado Borrador NO afecta los saldos ni los reportes hasta que sea contabilizado.')

doc.add_paragraph()
doc.add_paragraph('Ejemplo práctico — Registro de pago de servicios básicos:')

add_table(
    ['Cuenta', 'Descripción', 'Debe', 'Haber'],
    [
        ['5.1.03 - Servicios Básicos', 'Pago luz y agua enero', '$85.00', ''],
        ['1.1.01 - Caja', 'Pago en efectivo', '', '$85.00'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('Ejemplo práctico — Compra de mercadería al contado:')

add_table(
    ['Cuenta', 'Descripción', 'Debe', 'Haber'],
    [
        ['1.1.04 - Inventario', 'Compra de repuestos', '$200.00', ''],
        ['1.1.05 - IVA en Compras', 'IVA 15% de la compra', '$30.00', ''],
        ['1.1.01 - Caja', 'Pago al proveedor', '', '$230.00'],
    ]
)

# 4.4
doc.add_heading('4.4 Ver detalle de un asiento', level=2)
add_step(1, 'Ubicar', 'En el listado de asientos, encuentre el asiento que desea ver.')
add_step(2, 'Ver detalle', 'Haga clic en el ícono de "Ver" (ojo) en la columna de acciones.')
add_step(3, 'Revisar', 'Se abrirá una vista con toda la información del asiento: número, fecha, descripción, tipo, estado, y la tabla de líneas con las cuentas, descripciones y montos.')

# 4.5
doc.add_heading('4.5 Editar un asiento (solo borrador)', level=2)
add_step(1, 'Ubicar', 'Encuentre el asiento en estado "Borrador" que desea modificar.')
add_step(2, 'Editar', 'Haga clic en el ícono de "Editar" (lápiz) en la columna de acciones.')
add_step(3, 'Modificar', 'Puede cambiar la fecha, descripción y/o las líneas de detalle. Si modifica las líneas, se reemplazan todas las existentes.')
add_step(4, 'Guardar', 'Haga clic en "Guardar" para aplicar los cambios.')

add_note('Solo se pueden editar asientos en estado "Borrador". Los asientos contabilizados o anulados no se pueden modificar.')

# 4.6
doc.add_heading('4.6 Contabilizar un asiento', level=2)
doc.add_paragraph('Contabilizar un asiento significa aprobarlo oficialmente. Una vez contabilizado, el asiento afectará los saldos de las cuentas y aparecerá en todos los reportes.')

add_step(1, 'Ubicar', 'Encuentre el asiento en estado "Borrador".')
add_step(2, 'Contabilizar', 'Haga clic en el botón "Contabilizar" (ícono de check).')
add_step(3, 'Confirmar', 'Confirme la acción. El estado cambiará a "Contabilizado" (verde).')

add_note('Antes de contabilizar, revise cuidadosamente las líneas del asiento. Una vez contabilizado, no podrá editarlo. Si detecta un error después, deberá anular el asiento y crear uno nuevo.')

# 4.7
doc.add_heading('4.7 Anular un asiento', level=2)
doc.add_paragraph('Si un asiento tiene un error o no corresponde, puede anularlo:')

add_step(1, 'Ubicar', 'Encuentre el asiento que desea anular (puede estar en Borrador o Contabilizado).')
add_step(2, 'Anular', 'Haga clic en el botón "Anular" (ícono de X roja).')
add_step(3, 'Confirmar', 'Confirme la acción. El estado cambiará a "Anulado" (rojo).')

add_note('Un asiento anulado deja de afectar los saldos y no aparece en los reportes. El registro se mantiene para auditoría.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 5. GENERACIÓN AUTOMÁTICA
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Generación Automática de Asientos', level=1)

doc.add_paragraph(
    'Una de las funciones más útiles del módulo es la generación automática de asientos contables '
    'desde los documentos del sistema (facturas, gastos, retiros). Esto evita errores y ahorra tiempo.'
)

# 5.1
doc.add_heading('5.1 Desde una factura', level=2)

add_step(1, 'Ir al historial', 'Vaya al Historial de Facturas.')
add_step(2, 'Buscar factura', 'Ubique la factura para la cual desea generar el asiento.')
add_step(3, 'Generar', 'Haga clic en el botón "Generar Asiento Contable" en la fila de la factura.')

doc.add_paragraph('El sistema generará automáticamente el asiento según el tipo de venta:')

doc.add_paragraph()
p = add_paragraph('Venta de contado:', bold=True)

add_table(
    ['Cuenta', 'Debe', 'Haber'],
    [
        ['1.1.01 - Caja', 'Total factura', ''],
        ['4.1.01 - Ventas de Productos', '', 'Subtotal'],
        ['2.1.02 - IVA en Ventas', '', 'IVA'],
    ]
)

doc.add_paragraph()
p = add_paragraph('Venta a crédito:', bold=True)

add_table(
    ['Cuenta', 'Debe', 'Haber'],
    [
        ['1.1.03 - Cuentas por Cobrar', 'Total factura', ''],
        ['4.1.01 - Ventas de Productos', '', 'Subtotal'],
        ['2.1.02 - IVA en Ventas', '', 'IVA'],
    ]
)

doc.add_paragraph()
add_note('Si ya existe un asiento para esa factura, el sistema le mostrará un mensaje indicando el número del asiento existente. No se permiten asientos duplicados para la misma factura.')

# 5.2
doc.add_heading('5.2 Desde un gasto', level=2)

add_step(1, 'Ir a Gastos', 'Vaya al módulo de Gastos.')
add_step(2, 'Generar', 'Haga clic en "Generar Asiento" en la fila del gasto.')

doc.add_paragraph('El asiento generado será:')

add_table(
    ['Cuenta', 'Debe', 'Haber'],
    [
        ['5.1.07 - Gastos Generales', 'Monto del gasto', ''],
        ['1.1.01 - Caja', '', 'Monto del gasto'],
    ]
)

# 5.3
doc.add_heading('5.3 Desde un retiro de caja', level=2)

add_step(1, 'Ir a Retiros', 'Vaya a la sección de Retiros.')
add_step(2, 'Generar', 'Haga clic en "Generar Asiento" en la fila del retiro.')

doc.add_paragraph('El asiento generado será:')

add_table(
    ['Cuenta', 'Debe', 'Haber'],
    [
        ['5.1.07 - Gastos Generales', 'Valor del retiro', ''],
        ['1.1.01 - Caja', '', 'Valor del retiro'],
    ]
)

doc.add_paragraph()
add_tip('Los asientos generados automáticamente se crean en estado "Contabilizado" directamente, ya que provienen de transacciones ya aprobadas en el sistema.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 6. REPORTES CONTABLES
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Reportes Contables', level=1)

doc.add_paragraph(
    'Los reportes contables permiten analizar la situación financiera del negocio. '
    'Todos los reportes solo consideran asientos en estado "Contabilizado".'
)

# 6.1
doc.add_heading('6.1 Libro Diario', level=2)

doc.add_paragraph(
    'El Libro Diario muestra todos los asientos contables registrados en un período, '
    'en orden cronológico, con sus líneas de detalle.'
)

add_step(1, 'Seleccionar reporte', 'En la pestaña "Reportes", seleccione "Libro Diario".')
add_step(2, 'Definir período', 'Ingrese la fecha desde y fecha hasta.')
add_step(3, 'Generar', 'Haga clic en "Generar Reporte".')

doc.add_paragraph('El reporte mostrará cada asiento con sus líneas expandibles y al final los totales globales de Debe y Haber que deben ser iguales.')

p = add_paragraph('¿Para qué sirve?', bold=True)
doc.add_paragraph('Para revisar cronológicamente todas las transacciones del período y verificar que cada asiento esté correctamente registrado.')

# 6.2
doc.add_heading('6.2 Libro Mayor', level=2)

doc.add_paragraph(
    'El Libro Mayor muestra los movimientos agrupados por cuenta contable. '
    'Permite ver cuánto se movió en cada cuenta y su saldo resultante.'
)

add_step(1, 'Seleccionar reporte', 'Seleccione "Libro Mayor".')
add_step(2, 'Definir período', 'Ingrese la fecha desde y fecha hasta.')
add_step(3, 'Opcional — Filtrar por cuenta', 'Si desea ver los movimientos de una cuenta específica, selecciónela en el dropdown. Si no selecciona ninguna, verá el resumen de todas las cuentas.')
add_step(4, 'Generar', 'Haga clic en "Generar Reporte".')

doc.add_paragraph('Vista resumen (sin filtro de cuenta):')
doc.add_paragraph('Muestra una tabla con cada cuenta, su total Debe, total Haber y saldo. Puede hacer clic en una cuenta para ver sus movimientos individuales.')

doc.add_paragraph()
doc.add_paragraph('Vista detalle (con cuenta seleccionada):')
doc.add_paragraph('Muestra cada movimiento de la cuenta seleccionada con fecha, número de asiento, descripción, debe y haber.')

p = add_paragraph('¿Para qué sirve?', bold=True)
doc.add_paragraph('Para analizar el comportamiento de una cuenta específica. Por ejemplo, ver todos los movimientos de "Caja" para conciliar el efectivo.')

# 6.3
doc.add_heading('6.3 Balance de Comprobación', level=2)

doc.add_paragraph(
    'El Balance de Comprobación muestra las sumas y saldos de todas las cuentas con movimientos en el período. '
    'Es una herramienta de verificación: los totales de Saldo Debe y Saldo Haber deben ser iguales.'
)

add_step(1, 'Seleccionar reporte', 'Seleccione "Balance de Comprobación".')
add_step(2, 'Definir período', 'Ingrese la fecha desde y fecha hasta.')
add_step(3, 'Generar', 'Haga clic en "Generar Reporte".')

doc.add_paragraph('El reporte muestra para cada cuenta:')
add_table(
    ['Columna', 'Significado'],
    [
        ['Sumas Debe', 'Total de todos los débitos en el período'],
        ['Sumas Haber', 'Total de todos los créditos en el período'],
        ['Saldo Debe', 'Saldo a favor del debe (si corresponde)'],
        ['Saldo Haber', 'Saldo a favor del haber (si corresponde)'],
    ]
)

doc.add_paragraph()
p = add_paragraph('¿Para qué sirve?', bold=True)
doc.add_paragraph('Para verificar que la contabilidad esté cuadrada. Si los totales de Saldo Debe y Saldo Haber son iguales, la contabilidad está correcta.')

# 6.4
doc.add_heading('6.4 Balance General', level=2)

doc.add_paragraph(
    'El Balance General muestra la situación financiera de la empresa a una fecha de corte. '
    'Presenta los Activos por un lado y los Pasivos + Patrimonio por el otro.'
)

add_step(1, 'Seleccionar reporte', 'Seleccione "Balance General".')
add_step(2, 'Definir fecha de corte', 'Ingrese la fecha hasta la cual desea ver el balance.')
add_step(3, 'Generar', 'Haga clic en "Generar Reporte".')

doc.add_paragraph('El reporte se divide en:')
items = [
    'ACTIVOS: Todo lo que la empresa posee (caja, bancos, inventario, cuentas por cobrar, equipos).',
    'PASIVOS: Todo lo que la empresa debe (cuentas por pagar, IVA, préstamos).',
    'PATRIMONIO: Capital de los dueños y utilidades acumuladas.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = add_paragraph('Ecuación contable fundamental:', bold=True)
doc.add_paragraph('ACTIVOS = PASIVOS + PATRIMONIO')
doc.add_paragraph('Si esta ecuación no se cumple, hay un error en los registros contables.')

# 6.5
doc.add_heading('6.5 Estado de Resultados', level=2)

doc.add_paragraph(
    'El Estado de Resultados (también llamado Estado de Pérdidas y Ganancias) muestra '
    'los ingresos y gastos del período para determinar si hubo utilidad o pérdida.'
)

add_step(1, 'Seleccionar reporte', 'Seleccione "Estado de Resultados".')
add_step(2, 'Definir período', 'Ingrese la fecha desde y fecha hasta.')
add_step(3, 'Generar', 'Haga clic en "Generar Reporte".')

doc.add_paragraph('El reporte muestra:')
items = [
    'INGRESOS: Ventas de productos, servicios técnicos, otros ingresos.',
    'GASTOS: Costo de ventas, sueldos, servicios básicos, arriendo, gastos generales.',
    'UTILIDAD NETA: Ingresos - Gastos. Si es positivo, hay ganancia. Si es negativo, hay pérdida.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = add_paragraph('¿Para qué sirve?', bold=True)
doc.add_paragraph('Para conocer cuánto ganó o perdió el negocio en un período determinado. Es el reporte más importante para la toma de decisiones.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 7. PREGUNTAS FRECUENTES
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Preguntas Frecuentes', level=1)

faqs = [
    (
        '¿Puedo modificar un asiento que ya fue contabilizado?',
        'No. Los asientos contabilizados no se pueden editar para mantener la integridad contable. Si necesita corregir un error, debe anular el asiento y crear uno nuevo con los datos correctos.'
    ),
    (
        '¿Qué pasa si anulo un asiento contabilizado?',
        'El asiento deja de afectar los saldos de las cuentas y ya no aparece en los reportes. El registro se mantiene en el sistema con estado "Anulado" para fines de auditoría.'
    ),
    (
        '¿Puedo generar un asiento dos veces para la misma factura?',
        'No. El sistema valida que no exista un asiento activo (no anulado) para la misma factura. Si necesita regenerarlo, primero anule el asiento existente.'
    ),
    (
        '¿Los asientos en borrador afectan los reportes?',
        'No. Solo los asientos en estado "Contabilizado" afectan los saldos y aparecen en los reportes.'
    ),
    (
        '¿Puedo agregar más cuentas al plan de cuentas?',
        'Sí. Use el botón "Nueva Cuenta" en la pestaña Plan de Cuentas. Asegúrese de definir correctamente el tipo, naturaleza y cuenta padre.'
    ),
    (
        '¿Qué debo hacer al final de cada mes?',
        'Se recomienda: 1) Verificar que todos los asientos del mes estén contabilizados. 2) Generar el Balance de Comprobación para verificar que está cuadrado. 3) Generar el Estado de Resultados para conocer la utilidad del mes.'
    ),
    (
        '¿Qué significa que un asiento "no está cuadrado"?',
        'Significa que la suma del Debe no es igual a la suma del Haber. Revise las líneas del asiento y corrija los montos hasta que ambos totales sean iguales.'
    ),
    (
        '¿Quién puede ver el módulo de contabilidad?',
        'Solo los usuarios con tipo Administrador o Super Usuario. Los técnicos y personal de atención al público no tienen acceso a este módulo.'
    ),
]

for pregunta, respuesta in faqs:
    p = doc.add_paragraph()
    run = p.add_run(pregunta)
    run.bold = True
    doc.add_paragraph(respuesta)
    doc.add_paragraph()

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 8. GLOSARIO
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Glosario', level=1)

terms = [
    ('Activo', 'Bien o derecho que posee la empresa (caja, bancos, inventario, equipos, cuentas por cobrar).'),
    ('Asiento Contable', 'Registro de una transacción financiera que indica qué cuentas se afectan y por qué monto.'),
    ('Balance de Comprobación', 'Reporte que verifica que la contabilidad está cuadrada mostrando sumas y saldos de todas las cuentas.'),
    ('Balance General', 'Reporte que muestra la situación financiera (Activos = Pasivos + Patrimonio) a una fecha de corte.'),
    ('Contabilizar', 'Aprobar un asiento para que afecte los saldos de las cuentas y aparezca en los reportes.'),
    ('Debe', 'Lado izquierdo del asiento. Registra aumentos en activos/gastos y disminuciones en pasivos/patrimonio/ingresos.'),
    ('Estado de Resultados', 'Reporte que muestra ingresos menos gastos para determinar la utilidad o pérdida.'),
    ('Gasto', 'Erogación necesaria para la operación del negocio (sueldos, arriendo, servicios básicos).'),
    ('Haber', 'Lado derecho del asiento. Registra aumentos en pasivos/patrimonio/ingresos y disminuciones en activos/gastos.'),
    ('Ingreso', 'Dinero que entra a la empresa por ventas de productos o servicios.'),
    ('Libro Diario', 'Reporte cronológico de todos los asientos contables registrados.'),
    ('Libro Mayor', 'Reporte que agrupa los movimientos por cuenta contable.'),
    ('Naturaleza', 'Indica si una cuenta normalmente tiene saldo Deudor (activos, gastos) o Acreedor (pasivos, patrimonio, ingresos).'),
    ('Partida Doble', 'Principio contable que establece que todo Debe tiene un Haber. La suma de débitos siempre es igual a la suma de créditos.'),
    ('Pasivo', 'Deuda u obligación que tiene la empresa (cuentas por pagar, IVA, préstamos).'),
    ('Patrimonio', 'Capital de los dueños más las utilidades acumuladas del negocio.'),
    ('Plan de Cuentas', 'Estructura jerárquica de todas las cuentas contables de la empresa.'),
    ('Saldo', 'Diferencia entre el total de débitos y créditos de una cuenta.'),
    ('Utilidad Neta', 'Resultado de restar los gastos totales de los ingresos totales. Si es positivo hay ganancia, si es negativo hay pérdida.'),
]

for term, definition in terms:
    p = doc.add_paragraph()
    run = p.add_run(term + ': ')
    run.bold = True
    p.add_run(definition)


# ─── GUARDAR ───
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Manual_Modulo_Contabilidad.docx')
doc.save(output_path)
print(f'Documento generado: {output_path}')
